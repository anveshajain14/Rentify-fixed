# """
# chatbot.py

# High-level helpers for:
# - Intent detection
# - Policy QA (full-document grounding, no RAG)
# - Item discovery via MongoDB
# """

# import os
# import json
# import re
# from typing import List

# from pymongo import MongoClient
# from docx import Document
# import google.generativeai as genai
# from dotenv import load_dotenv

# # ------------------------------------------------------------------
# # Env & Gemini
# # ------------------------------------------------------------------

# load_dotenv()

# genai.configure(
#     api_key=os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
# )
# model = genai.GenerativeModel("gemini-2.5-flash")

# # ------------------------------------------------------------------
# # Mongo (read-only)
# # ------------------------------------------------------------------

# client = MongoClient(os.getenv("MONGODB_URI"))
# db = client["rentify"]
# items_collection = db["items"]

# # ------------------------------------------------------------------
# # Constants
# # ------------------------------------------------------------------

# POLICY_FALLBACK_RESPONSE = (
#     "I couldn’t find this information in our policy. "
#     "Please reach out to our customer support team for accurate assistance."
# )

# # ------------------------------------------------------------------
# # Generic helpers
# # ------------------------------------------------------------------

# def safe_json_parse(text: str) -> dict:
#     text = re.sub(r"```json|```", "", text).strip()
#     try:
#         return json.loads(text)
#     except Exception:
#         return {}

# # ------------------------------------------------------------------
# # Policy loading (FULL DOCUMENT, NO RAG)
# # ------------------------------------------------------------------

# def _load_policy_sections() -> List[str]:
#     """
#     Load policy document into logical sections.
#     Supports Policy.docx and policies.docx.
#     """
#     for path in ["Policy.docx", "policies.docx"]:
#         if os.path.exists(path):
#             doc = Document(path)
#             sections = []
#             current = ""

#             for para in doc.paragraphs:
#                 if para.text.strip():
#                     current += para.text.strip() + "\n"
#                 else:
#                     if current:
#                         sections.append(current.strip())
#                         current = ""

#             if current:
#                 sections.append(current.strip())

#             return sections

#     return []

# # Load once at startup (deterministic & stable)
# _POLICY_SECTIONS = _load_policy_sections()

# # ------------------------------------------------------------------
# # Intent detection
# # ------------------------------------------------------------------

# def detect_intent(query: str) -> str:
#     """
#     Returns: ITEM or POLICY
#     """
#     prompt = f"""
# Classify the user query into ONE category.
# Reply with exactly one word.

# ITEM   -> item search / recommendations
# POLICY -> rules, terms, deposits, late fees, etc.

# Query: {query}
# """
#     try:
#         label = model.generate_content(prompt).text.strip().upper()
#     except Exception:
#         return "POLICY"

#     if "ITEM" in label:
#         return "ITEM"
#     return "POLICY"

# # ------------------------------------------------------------------
# # Item discovery
# # ------------------------------------------------------------------

# def extract_item_filters(query: str) -> dict:
#     prompt = f"""
# Extract structured filters as valid JSON.

# Keys:
#   category  (string or null)
#   max_price (number or null)
#   location  (string or null)

# Query: {query}
# """
#     try:
#         return safe_json_parse(model.generate_content(prompt).text)
#     except Exception:
#         return {}

# def find_items(filters: dict):
#     q = {"availability": True}

#     if isinstance(filters.get("category"), str) and filters["category"]:
#         q["category"] = filters["category"]

#     if isinstance(filters.get("location"), str) and filters["location"]:
#         q["location"] = filters["location"]

#     # Read-only projection; internal IDs hidden
#     return list(items_collection.find(q, {"_id": 0}))

# # ------------------------------------------------------------------
# # Policy QA (e-commerce style fallback)
# # ------------------------------------------------------------------

# def answer_from_policy(query: str) -> str:
#     """
#     Answer policy questions using the FULL policy document.
#     Respond in a friendly, conversational e-commerce tone.
#     """
#     if not _POLICY_SECTIONS:
#         return POLICY_FALLBACK_RESPONSE

#     prompt = f"""
# You are a helpful rental marketplace chatbot.

# Answer the user's question based on the policy text below.
# DO NOT mention section numbers.
# DO NOT copy sentences verbatim.
# Paraphrase and explain the rules in a clear, friendly,
# customer-facing tone, like an e-commerce support assistant.

# If the information is not clearly available in the policy,
# reply EXACTLY:
# {POLICY_FALLBACK_RESPONSE}

# POLICY TEXT:
# {_POLICY_SECTIONS}

# QUESTION:
# {query}
# """
#     try:
#         answer = model.generate_content(prompt).text.strip()
#     except Exception:
#         return POLICY_FALLBACK_RESPONSE

#     if not answer or answer.strip().lower() == POLICY_FALLBACK_RESPONSE.lower():
#         return POLICY_FALLBACK_RESPONSE

#     return answer


#     if not answer or answer.strip().lower() == POLICY_FALLBACK_RESPONSE.lower():
#         return POLICY_FALLBACK_RESPONSE

#     return answer



"""
chatbot.py

High-level helpers for:
- Intent detection
- Policy QA (full-document grounding, no RAG)
- Item discovery via MongoDB
"""

import os
import json
import re
from typing import List

from pymongo import MongoClient
from docx import Document
from dotenv import load_dotenv
from groq import Groq

# ------------------------------------------------------------------
# Env & Groq Setup
# ------------------------------------------------------------------

load_dotenv()

groq_client = Groq(
    api_key=os.getenv("GROQ_API_KEY")
)

LLM_MODEL = "llama-3.1-8b-instant"

# ------------------------------------------------------------------
# Mongo (read-only)
# ------------------------------------------------------------------

client = MongoClient(os.getenv("MONGODB_URI"))
db = client["rentify"]
items_collection = db["products"]

# ------------------------------------------------------------------
# Constants
# ------------------------------------------------------------------

POLICY_FALLBACK_RESPONSE = (
    "I couldn’t find this information in our policy. "
    "Please reach out to our customer support team for accurate assistance."
)

# ------------------------------------------------------------------
# Generic helpers
# ------------------------------------------------------------------

def safe_json_parse(text: str) -> dict:
    text = re.sub(r"```json|```", "", text).strip()
    try:
        return json.loads(text)
    except Exception:
        return {}

def generate_llm_response(prompt: str) -> str:
    try:
        completion = groq_client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
        )
        return completion.choices[0].message.content.strip()
    except Exception:
        return ""

# ------------------------------------------------------------------
# Policy loading (FULL DOCUMENT, NO RAG)
# ------------------------------------------------------------------

def _load_policy_sections() -> List[str]:
    """
    Load policy document into logical sections.
    Supports Policy.docx and policies.docx.
    """
    for path in ["Policy.docx", "policies.docx", "/mnt/data/policies.docx"]:
        if os.path.exists(path):
            doc = Document(path)
            sections = []
            current = ""

            for para in doc.paragraphs:
                if para.text.strip():
                    current += para.text.strip() + "\n"
                else:
                    if current:
                        sections.append(current.strip())
                        current = ""

            if current:
                sections.append(current.strip())

            return sections

    return []

# Load once at startup
_POLICY_SECTIONS = _load_policy_sections()

# ------------------------------------------------------------------
# Intent detection
# ------------------------------------------------------------------

def detect_intent(query: str) -> str:
    """
    Returns: ITEM or POLICY
    """
    prompt = f"""
Classify the user query into ONE category.
Reply with exactly one word.

ITEM   -> item search / recommendations
POLICY -> rules, terms, deposits, late fees, etc.

Query: {query}
"""
    try:
        label = generate_llm_response(prompt).strip().upper()
    except Exception:
        return "POLICY"

    if "ITEM" in label:
        return "ITEM"
    return "POLICY"

# ------------------------------------------------------------------
# Item discovery
# ------------------------------------------------------------------

def extract_item_filters(query: str) -> dict:
    prompt = f"""
You are a strict JSON generator.

Extract filters from the user query.

Return ONLY valid JSON.
Do NOT explain.
Do NOT add extra text.
Do NOT use markdown.

Format exactly like this:
{{
  "category": string or null,
  "max_price": number or null,
  "location": string or null
}}

User Query:
{query}
"""
    try:
        response = generate_llm_response(prompt)
        # print("RAW FILTER RESPONSE:", response) 
        return safe_json_parse(response)
    except Exception:
        return {}


def find_items(filters: dict):
    q = {
        "isApproved": True
    }

    # Category filter (case-insensitive)
    if isinstance(filters.get("category"), str) and filters["category"]:
        q["category"] = {
            "$regex": f"^{filters['category']}$",
            "$options": "i"
        }

    # Price filter
    if isinstance(filters.get("max_price"), (int, float)):
        q["pricePerDay"] = {"$lte": filters["max_price"]}

    print("MONGO QUERY:", q)

    items = list(items_collection.find(q, {"_id": 0}))

    # Convert ObjectId fields to string (important!)
    for item in items:
        if "seller" in item:
            item["seller"] = str(item["seller"])

    return items


# ------------------------------------------------------------------
# Policy QA (e-commerce style fallback)
# ------------------------------------------------------------------

def answer_from_policy(query: str) -> str:
    """
    Answer policy questions using the FULL policy document.
    Respond in a friendly, conversational e-commerce tone.
    """
    if not _POLICY_SECTIONS:
        return POLICY_FALLBACK_RESPONSE

    prompt = f"""
You are a helpful rental marketplace chatbot.

Answer the user's question based on the policy text below.
DO NOT mention section numbers.
DO NOT copy sentences verbatim.
Paraphrase and explain the rules in a clear, friendly,
customer-facing tone, like an e-commerce support assistant.

If the information is not clearly available in the policy,
reply EXACTLY:
{POLICY_FALLBACK_RESPONSE}

POLICY TEXT:
{_POLICY_SECTIONS}

QUESTION:
{query}
"""
    try:
        answer = generate_llm_response(prompt)
    except Exception:
        return POLICY_FALLBACK_RESPONSE

    if not answer or answer.strip().lower() == POLICY_FALLBACK_RESPONSE.lower():
        return POLICY_FALLBACK_RESPONSE

    return answer